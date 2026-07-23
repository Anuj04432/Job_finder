"""
Step 1: Extract raw text from a resume file (PDF, DOCX, or image).

Supports:
- Native text PDFs (via pdfplumber)
- Scanned/image PDFs (via OCR fallback with pytesseract + PyMuPDF rendering)
- DOCX files (via python-docx, preserving paragraph and table content)
- Plain images (via pytesseract OCR)

Install dependencies:
    pip install pdfplumber PyMuPDF python-docx pytesseract pillow --break-system-packages

You also need the Tesseract binary installed on the system (not just the python package):
    Ubuntu/Debian: sudo apt-get install tesseract-ocr
    macOS:         brew install tesseract
    Windows:       https://github.com/UB-Mannheim/tesseract/wiki
"""

import os
import io
from pathlib import Path

import pdfplumber
import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image


# Minimum characters we'd expect from a "real" text-based page.
# If extraction falls below this, we assume it's a scanned/image page and OCR it instead.
MIN_CHARS_PER_PAGE = 40


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF. Tries direct text extraction first (fast, accurate
    for native PDFs). Falls back to OCR per-page if a page yields little/no text
    (common for scanned resumes exported as images).
    """
    text_chunks = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""

            if len(page_text.strip()) >= MIN_CHARS_PER_PAGE:
                text_chunks.append(page_text)
            else:
                # Fallback: render this page as an image and OCR it
                ocr_text = _ocr_pdf_page(file_path, page_num)
                text_chunks.append(ocr_text)

    return "\n\n".join(text_chunks).strip()


def _ocr_pdf_page(file_path: str, page_num: int, zoom: int = 3) -> str:
    """Render a single PDF page to an image and run OCR on it."""
    doc = fitz.open(file_path)
    page = doc.load_page(page_num)

    # Higher zoom = higher resolution = better OCR accuracy
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix)
    img_bytes = pix.tobytes("png")

    image = Image.open(io.BytesIO(img_bytes))
    doc.close()

    return pytesseract.image_to_string(image)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file, including paragraphs and table cells
    (resumes often use tables for layout, e.g. skills grids or two-column formats).
    """
    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_text_from_image(file_path: str) -> str:
    """Extract text from a plain image file (e.g. a photographed/scanned resume)."""
    image = Image.open(file_path)
    return pytesseract.image_to_string(image).strip()


def extract_resume_text(file_path: str) -> str:
    """
    Unified entry point: detects file type by extension and routes to the
    right extractor. Returns raw extracted text.
    """
    ext = Path(file_path).suffix.lower()

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"No such file: {file_path}")

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in (".docx",):
        text = extract_text_from_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        text = extract_text_from_image(file_path)
    elif ext == ".doc":
        raise ValueError(
            "Legacy .doc format is not directly supported. "
            "Convert to .docx first (e.g. with LibreOffice: "
            "`libreoffice --headless --convert-to docx file.doc`)."
        )
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not text or len(text.strip()) < 20:
        raise ValueError(
            "Extraction returned little or no text. The file may be corrupted, "
            "password-protected, or an unsupported scanned format."
        )

    return text


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python extract_resume_text.py <path_to_resume>")
        sys.exit(1)

    path = sys.argv[1]
    extracted = extract_resume_text(path)
    print(extracted)